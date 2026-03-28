import os
import re
from datetime import datetime

from docx import Document
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from config.settings import OUTPUTS_DIR


class Exporter:
    """Export meeting outputs to Markdown, PDF, and DOCX."""

    def __init__(
        self,
        summary: str,
        diarized_transcript: str,
        meeting_title: str,
        output_dir: str = OUTPUTS_DIR,
    ):
        self.summary = summary or ""
        self.diarized_transcript = diarized_transcript or ""
        self.meeting_title = (meeting_title or "Meeting Notes").strip()
        self.output_dir = output_dir
        self._timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

        try:
            os.makedirs(self.output_dir, exist_ok=True)
        except PermissionError:
            print(
                f"[Exporter] Permission denied while creating output directory: {self.output_dir}"
            )
        except OSError as exc:
            print(f"[Exporter] Could not prepare output directory '{self.output_dir}': {exc}")

    def export_markdown(self) -> str:
        """Write summary and transcript to a Markdown file and return its path."""
        file_path = self._build_output_path("md")

        content = (
            f"# {self.meeting_title}\n\n"
            "## Summary\n\n"
            f"{self.summary.strip()}\n\n"
            "## Transcript\n\n"
            f"{self.diarized_transcript.strip()}\n"
        )

        try:
            with open(file_path, "w", encoding="utf-8") as md_file:
                md_file.write(content)
            return file_path
        except PermissionError:
            print(f"[Exporter] Permission denied writing Markdown file: {file_path}")
            return ""
        except OSError as exc:
            print(f"[Exporter] Failed to write Markdown file '{file_path}': {exc}")
            return ""

    def export_pdf(self) -> str:
        """Create a styled PDF with title, summary, and transcript sections."""
        file_path = self._build_output_path("pdf")

        styles = getSampleStyleSheet()
        story = []

        story.append(Paragraph(self._escape_pdf_text(self.meeting_title), styles["Title"]))
        story.append(Spacer(1, 12))

        story.append(Paragraph("Summary", styles["Heading2"]))
        story.append(Spacer(1, 6))
        summary_text = self.summary.strip() or "(No summary provided)"
        for paragraph_text in self._split_paragraphs(summary_text):
            story.append(Paragraph(self._escape_pdf_text(paragraph_text), styles["BodyText"]))
            story.append(Spacer(1, 6))

        story.append(Spacer(1, 8))
        story.append(Paragraph("Transcript", styles["Heading2"]))
        story.append(Spacer(1, 6))
        transcript_text = self.diarized_transcript.strip() or "(No transcript provided)"
        for paragraph_text in self._split_paragraphs(transcript_text):
            story.append(Paragraph(self._escape_pdf_text(paragraph_text), styles["BodyText"]))
            story.append(Spacer(1, 4))

        try:
            doc = SimpleDocTemplate(file_path)
            doc.build(story)
            return file_path
        except PermissionError:
            print(f"[Exporter] Permission denied writing PDF file: {file_path}")
            return ""
        except OSError as exc:
            print(f"[Exporter] Failed to write PDF file '{file_path}': {exc}")
            return ""
        except Exception as exc:
            print(f"[Exporter] Failed to build PDF '{file_path}': {exc}")
            return ""

    def export_docx(self) -> str:
        """Create a DOCX document with heading and section structure."""
        file_path = self._build_output_path("docx")

        document = Document()
        document.add_heading(self.meeting_title, level=1)

        document.add_heading("Summary", level=2)
        summary_text = self.summary.strip() or "(No summary provided)"
        for paragraph_text in self._split_paragraphs(summary_text):
            document.add_paragraph(paragraph_text)

        document.add_heading("Transcript", level=2)
        transcript_text = self.diarized_transcript.strip() or "(No transcript provided)"
        for paragraph_text in self._split_paragraphs(transcript_text):
            document.add_paragraph(paragraph_text)

        try:
            document.save(file_path)
            return file_path
        except PermissionError:
            print(f"[Exporter] Permission denied writing DOCX file: {file_path}")
            return ""
        except OSError as exc:
            print(f"[Exporter] Failed to write DOCX file '{file_path}': {exc}")
            return ""

    def _build_output_path(self, extension: str) -> str:
        safe_title = self._safe_title(self.meeting_title)
        filename = f"{safe_title}_{self._timestamp}.{extension}"
        return os.path.join(self.output_dir, filename)

    def _safe_title(self, title: str) -> str:
        safe = re.sub(r"[^A-Za-z0-9_-]+", "_", title.strip())
        safe = re.sub(r"_+", "_", safe).strip("_")
        return safe or "meeting"

    def _split_paragraphs(self, text: str):
        lines = [line.strip() for line in text.splitlines()]
        paragraphs = [line for line in lines if line]
        return paragraphs if paragraphs else [""]

    def _escape_pdf_text(self, text: str) -> str:
        return (
            text.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
        )
